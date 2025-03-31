import random
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_problem(operation, max_digits_op1, max_digits_op2):
    max_val_op1 = 10**max_digits_op1 - 1
    max_val_op2 = 10**max_digits_op2 - 1
    a = random.randint(1, max_val_op1)
    b = random.randint(1, max_val_op2)

    if operation == "add":
        symbol = "+"
    elif operation == "subtract":
        symbol = "-"
        if a < b:
            a, b = b, a
    elif operation == "multiply":
        symbol = "×"
    elif operation == "divide":
        symbol = "÷"
        b = random.randint(1, max_val_op2)
        #        a = b * random.randint(1, max_val_op1 // b or 1)  # ensure clean division
        a = random.randint(b, max_val_op1)  # ensure clean division
    else:
        raise ValueError("Unsupported operation")

    return a, b, symbol

def format_problem(a, b, symbol):

    if symbol == "÷":
       # Long division style: divisor ) dividend with bar over the dividend
        dividend = str(a)
        divisor = str(b)
        bar = "_" * (len(dividend) + 2)  # Add a bit of padding
        line1 = f"{' ' * (len(divisor) + 2)}{bar}"
        line2 = f"{divisor} ) {dividend}"
        return f"{line1}\n{line2}"
    else:
        # Normal vertical format for +, -, ×
        a_str = str(a)
        b_str = str(b)
        width = max(len(a_str), len(b_str)) + 2
        top = a_str.rjust(width)
        mid = f"{symbol}{b_str.rjust(width - 1)}"
        line = "_" * width
        return f"{top}\n{mid}\n{line}"
#    a_str = str(a)
#    b_str = str(b)
#    width = max(len(a_str), len(b_str)) + 2  # +2 for padding

#    top = a_str.rjust(width)
#    mid = f"{symbol}{b_str.rjust(width - 1)}"
#    line = "_" * width

#    return f"{top}\n{mid}\n{line}"

def set_row_height(row, height_in_inches):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_in_inches * 2000)))  # 1 inch = 1440 twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def set_table_full_width(table, doc):
    # Get available width (page width minus margins)
    section = doc.sections[0]
    page_width = section.page_width
    margin_left = section.left_margin
    margin_right = section.right_margin
    usable_width = page_width - margin_left - margin_right

    # Set each column width equally
    col_width = usable_width // 2
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    # Remove padding from each cell
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for side in ("top", "start", "bottom", "end"):
                mar = OxmlElement(f"w:tcMar")
                mar_elem = OxmlElement(f"w:{side}")
                mar_elem.set(qn("w:w"), "0")
                mar_elem.set(qn("w:type"), "dxa")
                mar.append(mar_elem)
                tcPr.append(mar)


    
def create_math_worksheet(operation, max_digits_op1, max_digits_op2, filename="math_worksheet.docx"):
    
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.25)  # You can try 0.1 or 0 if needed

    # Add space after title
#    doc.add_paragraph("")

    # Generate problems
    problems = [format_problem(*generate_problem(operation,
                                                 max_digits_op1,
                                                 max_digits_op2)) for _ in range(6)]

    # Create a 3x2 table
    table = doc.add_table(rows=3, cols=2)
    set_table_full_width(table, doc)
    
    table.autofit = False

    for i, row in enumerate(table.rows):
        set_row_height(row, 2.3)  # About 2 inches tall per row (fits 3 rows/page)

        for j, cell in enumerate(row.cells):
            index = i * 2 + j
            if index < len(problems):
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Fix vertical spacing between lines
                para.paragraph_format.line_spacing = 1
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                
                run = para.add_run(problems[index])
                run.font.name = 'Courier New'
                run.font.size = Pt(18)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # Optional: set cell width explicitly
            cell.width = Inches(3)

    doc.save(filename)
    print(f"Worksheet saved to {filename}")

# Example usage:
create_math_worksheet(operation="divide", max_digits_op1=3, max_digits_op2=2)
